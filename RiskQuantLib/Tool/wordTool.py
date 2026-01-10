#!/usr/bin/python
#coding = utf-8

import pandas as pd
from docx import Document # use package python-docx

#<import>
#</import>

def replaceParagraphContent(paragraph,paraDict:dict):
    """
    This function will replace the paragraph content of word Document object.
    This replace is inplace, so this function returns None. This function is very
    like the str.format() function in python, it replaces any symbol like {symbolName}
    according to key->value relation in paraDict, if there is a {code} showed in
    word, and paraDict declares '{code}':'000001', then the content of {code} will
    be replaced to 000001.

    Parameters
    ----------
    paragraph : Object
        The python-docx paragraph object.
    paraDict : dict
        A python dict used to declare which symbol should be replaced and how should
        it be replaced.

    Returns
    -------
    None
    """
    text = paragraph.text.strip()
    replacePara = [i for i in paraDict.keys() if text.find(i) != -1]
    if len(replacePara) != 0:
        for i in replacePara:
            if type(paraDict[i]) is str:
                replaceTo = paraDict[i]
            elif hasattr(paraDict[i], 'strftime'):
                replaceTo = paraDict[i].strftime('%Y-%m-%d')
            else:
                replaceTo = str(paraDict[i])
            text = text.replace(i, replaceTo)
        text = text.replace('{', '').replace('}', '')
        stack = []
        for line in paragraph.runs:
            if len(stack) != 0 and line.text.find('}') == -1:
                stack.append(line.text)
                line.text = ''
            elif len(stack) != 0 and line.text.find('}') != -1:
                endIndex = line.text.index('}')
                stack.append(line.text[:endIndex])
                line.text = line.text[endIndex + 1:]
                paraToBeReplaced = '{' + ''.join(stack) + '}'
                line.text = paraDict[paraToBeReplaced] + line.text
                stack = []
                text = text[len(paraDict[paraToBeReplaced]):]

            elif line.text.find('{') == -1:
                text = text[len(line.text):]
            elif line.text.find('}') != -1 and len(stack) == 0:
                startIndex = line.text.index('{')
                endIndex = line.text.index('}')
                paraToBeReplaced = line.text[startIndex:endIndex + 1]
                line.text = line.text[:startIndex] + paraDict[paraToBeReplaced] + line.text[endIndex + 1:]
                text = text[len(paraDict[paraToBeReplaced]):]
            else:
                content = line.text.split('{')
                line.text = content[0]
                stack.append(content[-1])

def formatWordWithTemplate(originFilePath:str,targetFilePath:str,paraDict:dict):
    """
    This function will iterate across the whole docx file and check for each paragraph
    and table, if it finds any symbol like {symbolName}, it will replace it into the
    value declared by paraDict.

    Parameters
    ----------
    originFilePath : str
        The docx template file whose content you want to check and replace.
    targetFilePath : str
        The path you want to save the after-replacement docx file.
    paraDict : dict
        The key->value pairs that declare what the symbols should be replaced to.
        The key of this dict should be string type and like '{symbolName}', and
        the value of this dict should be string type.

    Returns
    -------
    None
    """
    document = Document(originFilePath)
    for num, paragraph in enumerate(document.paragraphs):
        replaceParagraphContent(paragraph,paraDict)
    for table in document.tables:
        for row in range(len(table.rows)):
            for col in range(len(table.columns)):
                for num,paragraph in enumerate(table.cell(row,col).paragraphs):
                    replaceParagraphContent(paragraph,paraDict)
    document.save(targetFilePath)

def formatTableWithTemplate(originFilePath:str, targetFilePath:str, tableIndex:int, dfInput:pd.DataFrame):
    """
    This function will change the table in docx file into the given dataframe.
    This change is inplace. A new file will be generated and replace the original
    docx file.
    
    You should make sure the shape of table in your docx file is equal or larger than your dfInput, including index and
    column, which means your table in docx should have shape of (n+1, k+1) while your dfInput has shape (n, k).

    Parameters
    ----------
    originFilePath : str
        The docx file path that you want to change one of its table into the given
        dataframe.
    targetFilePath : str
        The docx file path where you want to save the changed file.
    tableIndex : int
        The table index that showed which table you want to change.
    dfInput : pd.DataFrame
        The dataframe that the original table should be changed to.

    Returns
    -------
    None
    """
    document = Document(originFilePath)
    table = document.tables[tableIndex]
    for col in range(dfInput.shape[1]):
        cl = dfInput.columns[col]
        for num, paragraph in enumerate(table.cell(0, col+1).paragraphs):
            paragraph.text = cl if isinstance(cl, str) else format(round(cl, 2), ",") if isinstance(cl,float) else str(cl)
    for row in range(dfInput.shape[0]):
        idx = dfInput.index[row]
        for num,paragraph in enumerate(table.cell(row+1,0).paragraphs):
            paragraph.text = idx if isinstance(idx,str) else format(round(idx,2),",") if isinstance(idx,float) else str(idx)
        for col in range(dfInput.shape[1]):
            for num,paragraph in enumerate(table.cell(row+1,col+1).paragraphs):
                value = dfInput.iloc[row, col]
                paragraph.text = value if isinstance(value, str) else format(round(value, 2), ",") if isinstance(value,float) else str(value)
    document.save(targetFilePath)

#<wordTool>
#</wordTool>






